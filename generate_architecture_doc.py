"""
Generate Professional Architecture Document for Migration Analysis Tool
Creates a Word document with clear diagrams and formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_architecture_document():
    """Create comprehensive architecture document"""
    
    print("\n" + "="*60)
    print("  GENERATING ARCHITECTURE DOCUMENT")
    print("="*60 + "\n")
    
    # Create document
    doc = Document()
    
    # Configure default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    print("Creating title page...")
    title = doc.add_heading('Migration Analysis Tool', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Architecture & Flow Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph()
    
    info = doc.add_paragraph('Version 1.0\nBosch Engineering\nMay 19, 2026')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.runs[0]
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # Table of Contents
    print("Adding table of contents...")
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. System Overview',
        '2. High-Level Architecture',
        '3. Comparison Modes',
        '   3.1 Offline ↔ Offline',
        '   3.2 Online ↔ Online', 
        '   3.3 Online ↔ Offline (Hybrid)',
        '4. Feature Matrix',
        '5. Use Cases',
        '6. Limitations & Constraints',
        '7. Technical Stack'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.25) if item.startswith('   ') else Inches(0)
    
    doc.add_page_break()
    
    # 1. System Overview
    print("Adding system overview...")
    doc.add_heading('1. System Overview', 1)
    
    doc.add_heading('Purpose', 2)
    doc.add_paragraph(
        'The Migration Analysis Tool is a comprehensive solution for analyzing code '
        'migration differences between platforms, with integrated RTC/ALM support and '
        'optional AI-powered assistance.'
    )
    
    doc.add_heading('Key Capabilities', 2)
    capabilities = [
        'Multi-Mode Comparison: Three distinct comparison workflows',
        'RTC Integration: Direct snapshot fetching from IBM Rational Team Concert',
        'Comprehensive Reporting: HTML diffs, Excel reports, CSV exports',
        'AI Assistant: Optional local AI for analysis insights (Ollama-based)',
        'Standalone Deployment: Executable versions for team distribution'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. Architecture
    print("Adding architecture diagrams...")
    doc.add_heading('2. High-Level Architecture', 1)
    
    doc.add_heading('System Components', 2)
    
    arch_text = """
┌─────────────────────────────────────────────────────┐
│        MIGRATION ANALYSIS TOOL (Python 3.14)        │
└──────────────────┬──────────────────────────────────┘
                   │
       ┌───────────┼───────────┐
       │           │           │
       ▼           ▼           ▼
  ┌────────┐  ┌────────┐  ┌────────┐
  │  GUI   │  │ CORE   │  │   AI   │
  │ Layer  │  │ ENGINE │  │ (OPT)  │
  └────────┘  └────────┘  └────────┘
       │           │           │
       └───────────┼───────────┘
                   │
       ┌───────────┼───────────┐
       │           │           │
       ▼           ▼           ▼
  ┌─────────┐ ┌─────────┐ ┌─────────┐
  │Offline↔ │ │Online↔  │ │Online↔  │
  │Offline  │ │Online   │ │Offline  │
  └─────────┘ └─────────┘ └─────────┘
       │           │           │
       └───────────┼───────────┘
                   │
       ┌───────────┼───────────┐
       │           │           │
       ▼           ▼           ▼
  ┌─────────┐ ┌─────────┐ ┌─────────┐
  │   RTC   │ │  FILE   │ │ REPORT  │
  │  INTEG  │ │ ANALYSIS│ │  GEN    │
  └─────────┘ └─────────┘ └─────────┘
    """
    
    p = doc.add_paragraph(arch_text)
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    doc.add_page_break()
    
    # 3. Comparison Modes
    print("Adding comparison modes...")
    doc.add_heading('3. Comparison Modes', 1)
    
    # Mode 1: Offline ↔ Offline
    doc.add_heading('3.1 Offline ↔ Offline Mode', 2)
    doc.add_paragraph('Compare local folders or ZIP archives without network access.')
    
    doc.add_heading('Flow Diagram:', 3)
    offline_flow = """
User Inputs:
  [Platform Folder/ZIP] + [Project Folder/ZIP]
         │
         ▼
  [Extract ZIPs if needed]
         │
         ▼
  [Component Detection]
         │
         ▼
  [File Mapping (Auto/Manual)]
         │
         ▼
  [Compare Files]
  • Hash comparison
  • Diff generation
  • Comment detection
         │
         ▼
  [Generate Reports]
  • HTML diffs
  • Excel summary
  • CSV export
         │
         ▼
  [Display Results in GUI]
    """
    
    p = doc.add_paragraph(offline_flow)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    doc.add_heading('Use Cases:', 3)
    use_cases_offline = [
        'Compare local backups',
        'Validate releases before deployment',
        'Offline analysis without network',
        'Quick local verification'
    ]
    for uc in use_cases_offline:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph()
    
    # Mode 2: Online ↔ Online
    doc.add_heading('3.2 Online ↔ Online Mode', 2)
    doc.add_paragraph('Compare two RTC snapshots directly from the server.')
    
    doc.add_heading('Flow Diagram:', 3)
    online_flow = """
User Inputs:
  [RTC Snapshot URL 1] + [RTC Snapshot URL 2]
         │
         ▼
  [Authenticate to RTC (NTLM)]
         │
         ▼
  [Fetch Snapshot Metadata]
         │
         ▼
  [Component Selection Dialog]
         │
         ▼
  [Download Selected Components]
  (Parallel downloads)
         │
         ▼
  [Extract to Temp Folders]
         │
         ▼
  [Compare Files]
  (Same engine as offline)
         │
         ▼
  [Generate Reports]
         │
         ▼
  [Cleanup Temp Folders]
         │
         ▼
  [Display Results]
    """
    
    p = doc.add_paragraph(online_flow)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    doc.add_heading('Use Cases:', 3)
    use_cases_online = [
        'Compare RTC baselines',
        'Branch comparison before merge',
        'Release validation',
        'Historical analysis'
    ]
    for uc in use_cases_online:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_page_break()
    
    # Mode 3: Hybrid
    doc.add_heading('3.3 Online ↔ Offline (Hybrid) Mode', 2)
    doc.add_paragraph('Compare RTC snapshot with local workspace folder.')
    
    doc.add_heading('Flow Diagram:', 3)
    hybrid_flow = """
User Inputs:
  [RTC Snapshot URL] + [Local Folder Path]
         │
         ▼
  [Authenticate to RTC (if needed)]
         │
         ▼
  [Fetch RTC Snapshot]
         │
         ▼
  [Component Selection]
         │
         ▼
  [Download RTC Components]
         │
         ▼
  [Extract to Temp Folder]
         │
         ▼
  [Map with Local Folder]
         │
         ▼
  [Compare RTC vs Local]
         │
         ▼
  [Generate Reports]
         │
         ▼
  [Cleanup Temp Folder]
         │
         ▼
  [Display Results]
    """
    
    p = doc.add_paragraph(hybrid_flow)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    doc.add_heading('Use Cases:', 3)
    use_cases_hybrid = [
        'Validate local work against baseline',
        'Pre-commit verification',
        'Development progress tracking',
        'Sync verification'
    ]
    for uc in use_cases_hybrid:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Feature Matrix
    print("Adding feature matrix...")
    doc.add_heading('4. Feature Matrix', 1)
    
    doc.add_heading('Core Features Comparison', 2)
    
    # Create table
    table = doc.add_table(rows=14, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Feature'
    header_cells[1].text = 'Offline↔Offline'
    header_cells[2].text = 'Online↔Online'
    header_cells[3].text = 'Online↔Offline'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    features = [
        ['Local Folders', '✅', '❌', '✅'],
        ['ZIP Archives', '✅', '❌', '❌'],
        ['RTC Snapshots', '❌', '✅', '✅'],
        ['Component Detection', '✅', '✅', '✅'],
        ['Auto File Mapping', '✅', '✅', '✅'],
        ['Parallel Processing', '✅', '✅', '✅'],
        ['HTML Reports', '✅', '✅', '✅'],
        ['Excel Reports', '✅', '✅', '✅'],
        ['RTC Authentication', '❌', '✅', '✅'],
        ['NTLM Proxy Support', '❌', '✅', '✅'],
        ['Offline Operation', '✅', '❌', '⚠️'],
        ['Internet Required', '❌', '✅', '✅']
    ]
    
    for i, feature_row in enumerate(features):
        row = table.rows[i + 1]
        for j, value in enumerate(feature_row):
            row.cells[j].text = value
    
    doc.add_page_break()
    
    # 5. Limitations
    print("Adding limitations...")
    doc.add_heading('5. Limitations & Constraints', 1)
    
    doc.add_heading('Technical Limitations', 2)
    
    limitations = [
        ('File Size', 'Maximum 10MB per file for optimal performance'),
        ('Component Size', 'Up to 1GB per component'),
        ('Total Size', 'Recommended maximum 10GB total comparison'),
        ('File Count', 'Best performance with <50,000 files'),
        ('Network', 'Corporate proxy requires NTLM authentication'),
        ('RTC Access', 'VPN may be required for RTC connections'),
        ('Binary Files', 'Limited binary diff support'),
        ('Encoding', 'UTF-8 assumed for text files')
    ]
    
    table2 = doc.add_table(rows=len(limitations)+1, cols=2)
    table2.style = 'Light List Accent 1'
    
    # Headers
    table2.rows[0].cells[0].text = 'Item'
    table2.rows[0].cells[1].text = 'Constraint'
    
    for i, (item, constraint) in enumerate(limitations):
        table2.rows[i+1].cells[0].text = item
        table2.rows[i+1].cells[1].text = constraint
    
    doc.add_paragraph()
    
    doc.add_heading('AI Assistant Limitations', 2)
    ai_limits = [
        'Requires one-time 2GB model download',
        'Corporate network blocks download (use home/VPN)',
        'CPU-based inference (slower than cloud AI)',
        'Context limited to current comparison',
        'May hallucinate - verify important findings'
    ]
    for limit in ai_limits:
        doc.add_paragraph(limit, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. Technical Stack
    print("Adding technical stack...")
    doc.add_heading('6. Technical Stack', 1)
    
    doc.add_heading('Core Technologies', 2)
    tech_stack = [
        'Language: Python 3.14.4',
        'GUI Framework: tkinter',
        'Packaging: PyInstaller 6.20.0',
        'HTTP Client: requests, httpx',
        'Authentication: requests-ntlm, httpx-ntlm',
        'Reporting: openpyxl (Excel)',
        'Optional AI: Ollama 0.24.0 with llama3.2:3b'
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('System Requirements', 2)
    
    doc.add_heading('Minimum:', 3)
    min_req = [
        'OS: Windows 10 (64-bit)',
        'CPU: Dual-core 2.0 GHz',
        'RAM: 4 GB',
        'Disk: 1 GB free space',
        'Network: Required for RTC modes'
    ]
    for req in min_req:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('Recommended:', 3)
    rec_req = [
        'OS: Windows 11 (64-bit)',
        'CPU: Quad-core 3.0 GHz+',
        'RAM: 8 GB+',
        'Disk: 5 GB (for AI model)',
        'Network: 100+ Mbps'
    ]
    for req in rec_req:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_page_break()
    
    # 7. Deployment Options
    print("Adding deployment options...")
    doc.add_heading('7. Deployment Options', 1)
    
    deployments = [
        ('Python Source', 'Requires Python 3.8+, full flexibility', '23.7 KB'),
        ('Lightweight EXE', 'No dependencies, fast startup', '22.63 MB'),
        ('Full EXE', 'Complete standalone with SCM tools', '101.07 MB')
    ]
    
    table3 = doc.add_table(rows=len(deployments)+1, cols=3)
    table3.style = 'Medium Shading 1 Accent 1'
    
    table3.rows[0].cells[0].text = 'Option'
    table3.rows[0].cells[1].text = 'Description'
    table3.rows[0].cells[2].text = 'Size'
    
    for i, (option, desc, size) in enumerate(deployments):
        table3.rows[i+1].cells[0].text = option
        table3.rows[i+1].cells[1].text = desc
        table3.rows[i+1].cells[2].text = size
    
    # Save document
    output_path = 'C:\\Users\\WIW1COB\\WP-8152\\Migration_Analysis_Tool_Architecture.docx'
    print(f"\nSaving document to: {output_path}")
    doc.save(output_path)
    
    print("\n" + "="*60)
    print("  ✅ DOCUMENT CREATED SUCCESSFULLY!")
    print("="*60)
    print(f"\nFile: {output_path}")
    print(f"Size: {os.path.getsize(output_path) / 1024:.2f} KB")
    print("\nOpening document...")
    
    return output_path

if __name__ == '__main__':
    try:
        doc_path = create_architecture_document()
        
        # Open the document
        import subprocess
        subprocess.Popen(['start', '', doc_path], shell=True)
        
        print("\n✅ Document opened in Microsoft Word!")
        print("\nYou can now:")
        print("  • Review the architecture diagrams")
        print("  • Add company branding")
        print("  • Customize formatting")
        print("  • Export as PDF")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("\nTroubleshooting:")
        print("  1. Ensure python-docx is installed: pip install python-docx")
        print("  2. Check write permissions in the folder")
        print("  3. Close any open Word documents")
